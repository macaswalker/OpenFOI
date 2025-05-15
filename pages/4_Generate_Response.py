import streamlit as st
import base64
from openfoi.config import FOI_EXEMPTIONS, RESPONSES_DIR
from openfoi.redaction import apply_redactions
from openfoi.data import save_requests
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os
import re
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

if "current_request_id" not in st.session_state:
    st.error("No request in session.")
    st.stop()
req_id = st.session_state.current_request_id
req = st.session_state.requests[req_id]

st.title("Generate FOI Response")

# 1. Apply redactions to markdown
redacted = apply_redactions(st.session_state.response_text, st.session_state.redactions)
letter = f"""# Response to your FOI request {req['reference']}

## Request

{req['request_text']}

## Response

{redacted}
"""  # Replace with better template if needed

# 2. Allow editing before export
letter_edit = st.text_area("Letter (Markdown)", letter, height=400)

def to_latin1(text):
    return text.encode('latin-1', 'replace').decode('latin-1')

def _to_pdf(text: str, path, blackout=False):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    if blackout:
        for line in text.split("\n"):
            parts = re.split(r"\*\*\[REDACTED ([A-Z0-9]+)\]\*\*", line)
            for part in parts:
                if part in FOI_EXEMPTIONS:
                    box_width = 40
                    box_height = 10
                    pdf.set_fill_color(0, 0, 0)
                    pdf.cell(box_width, box_height, f" [{part}] ", fill=True)
                else:
                    pdf.cell(0, 10, part, ln=0)
            pdf.ln(10)
    else:
        pdf.multi_cell(0, 10, text)
    pdf.output(str(path))

def _markdown_to_docx(md_text, doc, req):
    """
    Convert markdown-like text (with # for headers, ** for bold, [REDACTED Sxx]) into
    a nicely formatted docx with colored redactions.
    """
    lines = md_text.splitlines()
    for i, line in enumerate(lines):
        # Header conversion
        if line.startswith("### "):
            p = doc.add_paragraph(line[4:])
            p.style = 'Heading 3'
        elif line.startswith("## "):
            p = doc.add_paragraph(line[3:])
            p.style = 'Heading 2'
        elif line.startswith("# "):
            p = doc.add_paragraph(line[2:])
            p.style = 'Heading 1'
        else:
            # Handle redactions (bold red text) and bold (**text**)
            p = doc.add_paragraph()
            cursor = 0
            # Regex: **[REDACTED Sxx]**
            for match in re.finditer(r"\*\*\[REDACTED ([A-Z0-9]+)\]\*\*", line):
                # Text before
                if match.start() > cursor:
                    normal_run = p.add_run(line[cursor:match.start()])
                    normal_run.font.size = Pt(12)
                # The redacted bit
                code = match.group(1)
                red_run = p.add_run(f"[REDACTED {code}]")
                red_run.bold = True
                red_run.font.color.rgb = RGBColor(200, 0, 0)  # Bright red
                red_run.font.size = Pt(12)
                cursor = match.end()
            # Any remaining text after last redaction
            if cursor < len(line):
                normal_run = p.add_run(line[cursor:])
                normal_run.font.size = Pt(12)
    # Optionally add footer
    doc.add_paragraph("\n---\nThis is a response under the Freedom of Information Act 2000.")


def strip_markdown(text):
    """
    Removes markdown formatting, keeping only content and custom redacted markers.
    - Removes bold (**) and headers (#, ##, ###) from the start of lines.
    """
    # Remove bold markers
    text = text.replace("**", "")
    # Remove headers (replace starting #, ##, ### with just a blank)
    text = re.sub(r"^#{1,3}\s*", "", text, flags=re.MULTILINE)
    # Remove any stray markdown artifacts (add more here if needed)
    return text

def _markdown_to_docx(md_text, doc):
    """
    Convert stripped markdown-like text with [REDACTED Sxx] markers into docx, making redactions bold red.
    """
    lines = md_text.splitlines()
    for line in lines:
        p = doc.add_paragraph()
        cursor = 0
        for match in re.finditer(r"\[REDACTED ([A-Z0-9]+)\]", line):
            # Text before the redacted
            if match.start() > cursor:
                normal_run = p.add_run(line[cursor:match.start()])
                normal_run.font.size = Pt(12)
            # The redacted bit
            code = match.group(1)
            red_run = p.add_run(f"[REDACTED {code}]")
            red_run.bold = True
            red_run.font.color.rgb = RGBColor(200, 0, 0)  # Bright red
            red_run.font.size = Pt(12)
            cursor = match.end()
        # Any text after last redaction
        if cursor < len(line):
            normal_run = p.add_run(line[cursor:])
            normal_run.font.size = Pt(12)

def _to_docx(md_text: str, path, req):
    doc = Document()
    # Title
    title = f"Freedom of Information Act 2000 – Response [{req['reference']}]"
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date/reference block
    meta = doc.add_paragraph()
    meta.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    meta_run = meta.add_run(f"Reference: {req['reference']}\nDate: {datetime.now().strftime('%d %B %Y')}")
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph("")  # Spacer

    # Main body
    clean_text = strip_markdown(md_text)
    _markdown_to_docx(clean_text, doc)

    # Footer
    doc.add_paragraph("\n---\nThis is a response under the Freedom of Information Act 2000.")

    doc.save(str(path))


if "generated_files" not in st.session_state:
    st.session_state.generated_files = {}

if st.button("Generate Files (PDF & DOCX)"):
    now = datetime.now().strftime("%Y%m%d_%H%M%S")
    resp_pdf = RESPONSES_DIR / f"response_{req['reference']}_{now}.pdf"
    red_pdf  = RESPONSES_DIR / f"redacted_{req['reference']}_{now}.pdf"
    resp_docx = RESPONSES_DIR / f"response_{req['reference']}_{now}.docx"
    red_docx  = RESPONSES_DIR / f"redacted_{req['reference']}_{now}.docx"

    _to_pdf(to_latin1(letter_edit), resp_pdf)
    _to_pdf(to_latin1(redacted), red_pdf, blackout=True)
    _to_docx(letter_edit, resp_docx, req)
    _to_docx(redacted, red_docx, req)

    st.session_state.generated_files = {
        "Letter (PDF)": resp_pdf,
        "Redacted (PDF)": red_pdf,
        "Letter (DOCX)": resp_docx,
        "Redacted (DOCX)": red_docx,
    }
    req["status"] = "Completed"
    save_requests(st.session_state.requests)
    st.success("Files generated and request marked completed.")

# Always show download links if files are present in session
if st.session_state.get("generated_files"):
    for label, path in st.session_state.generated_files.items():
        if os.path.exists(path):
            b64 = base64.b64encode(path.read_bytes()).decode()
            if label.endswith(".PDF"):
                mime = "application/pdf"
            else:
                mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            href = f'<a class="download-btn" href="data:{mime};base64,{b64}" download="{path.name}">Download {label}</a>'
            st.markdown(href, unsafe_allow_html=True)
        else:
            st.error(f"{label} file was not created: {path}")
